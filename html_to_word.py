from flask import Flask, request, send_file
from docx import Document
from docx.shared import Pt

app = Flask(__name__)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        # Get paper content from the form
        paper_content = request.form['paper_content']

        # Create a Word document
        doc = Document()


        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal'].font.size = Pt(12)


        paragraph = doc.add_paragraph()
        run = paragraph.add_run(paper_content)


        run.font.name = '宋体'


        temp_file = 'temp.docx'
        doc.save(temp_file)


        return send_file(temp_file, as_attachment=True, download_name='paper.docx')


    return '''
        <form method="post">
            <textarea name="paper_content" rows="10" cols="50"></textarea><br>
            <input type="submit" value="Export to Word">
        </form>
    '''

if __name__ == '__main__':
    app.run(debug=True)
